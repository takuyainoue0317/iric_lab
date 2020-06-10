# -*- coding: utf-8 -*-
"""
Created on Mon Mar  9 19:28:27 2020

@author: riverlink
"""
#from attrdict import AttrDict

import os
import glob

import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

#================================================================================
# main
#================================================================================
def main():

    #デフォルト
    img_dir = "./img"
    f = "./result.docx"
    
    #テンプレートをコピー
    if os.path.isfile("./template.docx"):
        shutil.copyfile("template.docx", f)
        
    else:
        exit()
    
    #コピーしたファイルを開く
    doc = Document(f)

    #図のサイズ
    a4_h = 0.1 * (297 - (35 + 30))      #mm
    #a4_w = 0.1 * (210 - 30 * 2)         #mm
    #inch2mm = 25.4               #mm
    h = 0.85 * a4_h / 2.0
    #w = 0.9 * a4_w / 2.0
    
    #図を貼り付ける
    img_list = glob.glob(img_dir + "/*.png")
    count = 1
    for ff in img_list:

        #テーブルを追加
        table = doc.add_table(rows=2, cols=1)
        table.autofit=False
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        #図
        r1 = table.rows[0]
        pp = r1.cells[0].add_paragraph()
        pp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        r = pp.add_run()        
        r.add_picture(ff, height=Cm(h))
        #r.add_picture(ff)
        
        #キャプション
        r2 = table.rows[1]
        r2.cells[0].text = '図 - ' + str(count)
        r2.cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        r2.cells[0].paragraphs[0].runs[0].font.bold = True
        
        
#        #    doc.add_table(2,1)
#        
#        doc.add_picture(ff, height=Cm(h))
#        p = doc.paragraphs[-1]
#        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#        
#        p = doc.add_paragraph(" 図 - " + str(count))
#        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        #doc.add_picture(ff, width=Inches(w), height=Inches(h))
        #doc.add_page_break()
        count = count + 1
    doc.save(f)
    

#================================================================================
# main
#================================================================================
if __name__=='__main__': 

    main()
    
    